import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageTk


class LetterGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Генератор официальных писем")
        self.root.geometry("1000x700")
        self.root.minsize(900, 600)

        self.logo_path = None
        self.stamp_path = None
        self.logo_img = None
        self.stamp_img = None
        self.attachments = []

        self.create_widgets()

    def create_widgets(self):
        main_container = ttk.Frame(self.root)
        main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        input_frame = ttk.LabelFrame(main_container, text="Параметры письма")
        input_frame.pack(side=tk.LEFT, fill=tk.Y, padx=5, pady=5)

        fields = [
            ("Название компании:", "sender_company"),
            ("ИНН:", "inn"),
            ("КПП:", "kpp"),
            ("ОГРН:", "ogrn"),
            ("Юридический адрес:", "legal_address"),
            ("Почтовый адрес:", "post_address"),
            ("Телефон/факс:", "phone"),
            ("Исх. №:", "outgoing_number"),
            ("Дата исх.:", "outgoing_date"),
            ("Должность отправителя:", "sender_position"),
            ("ФИО отправителя:", "sender_name"),
            ("Компания получателя:", "recipient_company"),
            ("Должность получателя:", "recipient_position"),
            ("ФИО получателя:", "recipient_name"),
        ]

        self.entries = {}
        for i, (label, key) in enumerate(fields):
            ttk.Label(input_frame, text=label).grid(
                row=i, column=0, sticky="w", padx=5, pady=2
            )
            entry = ttk.Entry(input_frame, width=35)
            entry.grid(row=i, column=1, sticky="ew", padx=5, pady=2)
            entry.bind("<KeyRelease>", lambda e: self.update_preview())
            self.entries[key] = entry

        ttk.Label(input_frame, text="Текст письма:").grid(
            row=len(fields), column=0, columnspan=2, sticky="w", pady=5
        )
        self.letter_body = tk.Text(input_frame, height=10, width=40, wrap=tk.WORD)
        self.letter_body.grid(
            row=len(fields) + 1, column=0, columnspan=2, sticky="ew", padx=5, pady=5
        )
        self.letter_body.bind("<KeyRelease>", lambda e: self.update_preview())
        self.letter_body.bind("<Control-v>", self.paste_text)
        self.letter_body.bind("<Control-V>", self.paste_text)
        self.letter_body.bind("<<Paste>>", self.paste_text)

        btn_frame = ttk.Frame(input_frame)
        btn_frame.grid(row=len(fields) + 2, column=0, columnspan=2, pady=10)
        ttk.Button(btn_frame, text="Загрузить логотип", command=self.load_logo).pack(
            side=tk.LEFT, padx=5
        )
        ttk.Button(btn_frame, text="Загрузить печать", command=self.load_stamp).pack(
            side=tk.LEFT, padx=5
        )
        ttk.Button(
            btn_frame, text="Добавить вложение", command=self.add_attachment
        ).pack(side=tk.LEFT, padx=5)
        self.load_text_btn = ttk.Button(
            btn_frame, text="Загрузить текст", command=self.load_text_from_file
        )
        self.load_text_btn.pack(side=tk.LEFT, padx=5)

        preview_frame = ttk.LabelFrame(main_container, text="Предпросмотр")
        preview_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.preview_text = tk.Text(
            preview_frame,
            wrap=tk.WORD,
            font=("Courier New", 10),
            padx=10,
            pady=10,
            height=30,
            width=80,
        )
        self.preview_text.pack(fill=tk.BOTH, expand=True)

        scrollbar = ttk.Scrollbar(preview_frame, command=self.preview_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.preview_text.config(yscrollcommand=scrollbar.set)

        save_btn = ttk.Button(
            main_container, text="Сохранить в Word", command=self.save_document
        )
        save_btn.pack(side=tk.BOTTOM, pady=10, anchor=tk.SE, padx=20)

    def load_text_from_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt")])
        if file_path:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="cp1251") as f:
                    content = f.read()
            self.letter_body.delete("1.0", tk.END)
            self.letter_body.insert("1.0", content)
            self.update_preview()

    def paste_text(self, event=None):
        try:
            self.letter_body.insert(tk.INSERT, self.root.clipboard_get())
        except tk.TclError:
            pass
        return "break"

    def add_attachment(self):
        """Добавить вложение"""
        attachment_window = tk.Toplevel(self.root)
        attachment_window.title("Добавить вложение")
        attachment_window.geometry("600x400")

        attachment_text = tk.Text(attachment_window, wrap=tk.WORD, height=15, width=50)
        attachment_text.pack(padx=10, pady=10)

        def save_attachment():
            attachment_content = attachment_text.get("1.0", tk.END).strip()
            if attachment_content:
                self.attachments.append(attachment_content)
                attachment_window.destroy()

        save_button = ttk.Button(
            attachment_window, text="Сохранить вложение", command=save_attachment
        )
        save_button.pack(pady=5)

    def load_logo(self):
        self.logo_path = filedialog.askopenfilename(
            filetypes=[("Image Files", "*.png;*.jpg;*.jpeg")]
        )
        if self.logo_path:
            try:
                self.logo_img = Image.open(self.logo_path)
                self.logo_img.thumbnail((100, 100))
                self.logo_img = ImageTk.PhotoImage(self.logo_img)
                self.update_preview()
            except Exception as e:
                messagebox.showerror(
                    "Ошибка", f"Не удалось загрузить логотип: {str(e)}"
                )

    def load_stamp(self):
        self.stamp_path = filedialog.askopenfilename(
            filetypes=[("Image Files", "*.png;*.jpg;*.jpeg")]
        )
        if self.stamp_path:
            try:
                self.stamp_img = Image.open(self.stamp_path)
                self.stamp_img.thumbnail((100, 100))
                self.stamp_img = ImageTk.PhotoImage(self.stamp_img)
                self.update_preview()
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось загрузить печать: {str(e)}")

    def generate_content(self):
        try:
            data = {key: entry.get() for key, entry in self.entries.items()}
            letter_body = self.letter_body.get("1.0", tk.END).strip()

            recipient_name_parts = data["recipient_name"].split()
            salutation = (
                f"Уважаемый {recipient_name_parts[0]} {recipient_name_parts[1]}!"
                if len(recipient_name_parts) >= 2
                else "Уважаемый получатель!"
            )

            content = [
                f"{data['sender_company']}",
                f"ИНН {data['inn']} КПП {data['kpp']}",
                f"ОГРН {data['ogrn']}",
                f"{data['legal_address']}",
                f"{data['post_address']}",
                f"Тел./факс: {data['phone']}\n",
                f"{data['recipient_position']}",
                f"{data['recipient_company']}",
                f"{data['recipient_name']}\n",
                f"Исх. № {data['outgoing_number']} от {data['outgoing_date']}\n",
                f"{salutation}\n\n",
                letter_body + "\n\n",
                "С уважением,",
                f"{data['sender_position']}",
                f"{data['sender_company']}",
                f"{data['sender_name']}",
            ]
            return "\n".join(content)
        except Exception as e:
            return f"Ошибка формирования: {str(e)}"

    def update_preview(self, event=None):
        self.preview_text.config(state=tk.NORMAL)
        self.preview_text.delete(1.0, tk.END)

        if self.logo_img:
            self.preview_text.image_create(tk.END, image=self.logo_img)
            self.preview_text.insert(tk.END, "\n\n")

        content = self.generate_content()
        self.preview_text.insert(tk.END, content)

        if self.stamp_img:
            self.preview_text.insert(tk.END, "\n\n")
            self.preview_text.image_create(tk.END, image=self.stamp_img)

        self.preview_text.config(state=tk.DISABLED)

    def set_preview_styles(self):
        self.preview_text.tag_configure(
            "bold", font=("Times New Roman", 14, "bold"), justify="center"
        )
        self.preview_text.tag_configure("center", justify="center")
        self.preview_text.tag_configure(
            "normal", font=("Times New Roman", 12), justify="left"
        )
        self.preview_text.tag_configure(
            "justified", font=("Times New Roman", 12), justify="justify"
        )
        self.preview_text.tag_configure("red", foreground="red")
        self.preview_text.tag_configure("right", justify="right")

    def save_document(self):
        try:
            doc = Document()
            self.set_document_styles(doc)

            if self.logo_path:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(self.logo_path, width=Inches(1.5))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            company_name = doc.add_paragraph()
            company_name_run = company_name.add_run(
                self.entries["sender_company"].get()
            )
            company_name_run.bold = True
            company_name_run.font.size = Pt(14)
            company_name.alignment = WD_ALIGN_PARAGRAPH.CENTER

            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(3.5)
            table.columns[1].width = Inches(3.5)

            sender_cell, recipient_cell = table.rows[0].cells
            sender_text = f"ИНН {self.entries['inn'].get()} КПП {self.entries['kpp'].get()}\nОГРН {self.entries['ogrn'].get()}\n{self.entries['legal_address'].get()}\n{self.entries['post_address'].get()}\nТел./факс: {self.entries['phone'].get()}"
            recipient_text = f"{self.entries['recipient_position'].get()}\n{self.entries['recipient_company'].get()}\n{self.entries['recipient_name'].get()}"
            sender_cell.text = sender_text
            recipient_cell.text = recipient_text
            sender_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            recipient_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            outgoing = doc.add_paragraph()
            outgoing.add_run(
                f"Исх. № {self.entries['outgoing_number'].get()} от {self.entries['outgoing_date'].get()}"
            )
            outgoing.alignment = WD_ALIGN_PARAGRAPH.LEFT

            recipient_name = self.entries["recipient_name"].get().split()
            salutation = (
                "Уважаемый получатель!"
                if len(recipient_name) < 2
                else f"Уважаемый {recipient_name[0]} {recipient_name[1]}!"
            )
            doc.add_paragraph(salutation).paragraph_format.space_after = Pt(12)

            body_text = self.letter_body.get("1.0", tk.END).strip()
            for paragraph in body_text.split("\n"):
                p = doc.add_paragraph()
                p.add_run(f"  {paragraph.strip()}")
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Inches(0.5)
                p.paragraph_format.line_spacing = 1.5

            sign_block = [
                "С уважением,",
                self.entries["sender_position"].get(),
                self.entries["sender_company"].get(),
                self.entries["sender_name"].get(),
            ]
            for line in sign_block:
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if self.stamp_path:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(self.stamp_path, width=Inches(1.5))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx", filetypes=[("Word Documents", "*.docx")]
            )

            if self.attachments:
                doc.add_page_break()
                doc.add_paragraph("Вложения:\n")
                for attachment in self.attachments:
                    doc.add_paragraph(attachment)

            if file_path:
                doc.save(file_path)
                messagebox.showinfo("Успех", "Документ успешно сохранен!")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить документ: {str(e)}")

    def set_document_styles(self, doc):
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(0)

        section = doc.sections[0]
        section.left_margin = Inches(0.7874)
        section.right_margin = Inches(0.3937)
        section.top_margin = Inches(0.7874)
        section.bottom_margin = Inches(0.7874)


if __name__ == "__main__":
    root = tk.Tk()
    app = LetterGeneratorApp(root)
    root.mainloop()
